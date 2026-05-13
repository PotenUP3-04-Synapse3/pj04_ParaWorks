import io
from pypdf import PdfReader
from docx import Document as DocxDocument
from backend.app.documents.parsers import (
    DocumentParser,
    ParsedDocument,
    ParsedDocumentChunk,
    ParserRun,
    DocumentParserError,
)

class PdfDocumentParser:
    parser_name = 'pypdf'

    def parse(self, payload: bytes, *, metadata: dict[str, object]) -> ParsedDocument:
        try:
            reader = PdfReader(io.BytesIO(payload))
            chunks: list[ParsedDocumentChunk] = []
            chunk_index = 0
            for page_num, page in enumerate(reader.pages):
                text = page.extract_text()
                if text:
                    text = text.strip()
                    if text:
                        chunks.append(ParsedDocumentChunk(
                            chunk_index=chunk_index,
                            text=f"[Page {page_num + 1}]\n{text}",
                            source_snippet=text[:240].replace('\n', ' '),
                            page_number=page_num + 1,
                        ))
                        chunk_index += 1

            return _build_parsed_document(
                metadata=metadata,
                parser_name=self.parser_name,
                parser_status='parsed',
                parser_status_reason=None,
                chunks=chunks
            )
        except Exception as e:
            return _build_parsed_document(
                metadata=metadata,
                parser_name=self.parser_name,
                parser_status='error',
                parser_status_reason=f'Failed to parse PDF: {str(e)}',
                chunks=[]
            )

class DocxDocumentParser:
    parser_name = 'python-docx'

    def parse(self, payload: bytes, *, metadata: dict[str, object]) -> ParsedDocument:
        try:
            doc = DocxDocument(io.BytesIO(payload))
            chunks: list[ParsedDocumentChunk] = []
            chunk_index = 0
            
            # Simple chunking by paragraph (or grouping)
            current_chunk_text = []
            current_length = 0
            MAX_LENGTH = 1200
            
            def add_chunk():
                nonlocal chunk_index, current_chunk_text, current_length
                if current_chunk_text:
                    text = '\n'.join(current_chunk_text)
                    chunks.append(ParsedDocumentChunk(
                        chunk_index=chunk_index,
                        text=text,
                        source_snippet=text[:240].replace('\n', ' '),
                    ))
                    chunk_index += 1
                    current_chunk_text = []
                    current_length = 0

            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                if current_length + len(text) > MAX_LENGTH and current_length > 0:
                    add_chunk()
                current_chunk_text.append(text)
                current_length += len(text)
            add_chunk()

            return _build_parsed_document(
                metadata=metadata,
                parser_name=self.parser_name,
                parser_status='parsed',
                parser_status_reason=None,
                chunks=chunks
            )
        except Exception as e:
            return _build_parsed_document(
                metadata=metadata,
                parser_name=self.parser_name,
                parser_status='error',
                parser_status_reason=f'Failed to parse DOCX: {str(e)}',
                chunks=[]
            )

class TextDocumentParser:
    parser_name = 'built-in-text'

    def parse(self, payload: bytes, *, metadata: dict[str, object]) -> ParsedDocument:
        try:
            # Detect or assume UTF-8 with fallback
            try:
                text = payload.decode('utf-8')
            except UnicodeDecodeError:
                text = payload.decode('cp949', errors='replace') # Common Korean encoding fallback

            chunks: list[ParsedDocumentChunk] = []
            chunk_index = 0
            
            # Simple chunking by paragraph (double newline)
            paragraphs = [p.strip() for p in text.replace('\r\n', '\n').split('\n\n') if p.strip()]
            
            current_chunk_text = []
            current_length = 0
            MAX_LENGTH = 1200
            
            def add_chunk():
                nonlocal chunk_index, current_chunk_text, current_length
                if current_chunk_text:
                    full_text = '\n\n'.join(current_chunk_text)
                    chunks.append(ParsedDocumentChunk(
                        chunk_index=chunk_index,
                        text=full_text,
                        source_snippet=full_text[:240].replace('\n', ' '),
                    ))
                    chunk_index += 1
                    current_chunk_text = []
                    current_length = 0

            for para in paragraphs:
                if not para:
                    continue
                if current_length + len(para) > MAX_LENGTH and current_length > 0:
                    add_chunk()
                current_chunk_text.append(para)
                current_length += len(para)
            add_chunk()

            return _build_parsed_document(
                metadata=metadata,
                parser_name=self.parser_name,
                parser_status='parsed',
                parser_status_reason=None,
                chunks=chunks
            )
        except Exception as e:
            return _build_parsed_document(
                metadata=metadata,
                parser_name=self.parser_name,
                parser_status='error',
                parser_status_reason=f'Failed to parse Text/MD: {str(e)}',
                chunks=[]
            )

def _build_parsed_document(
    metadata: dict[str, object],
    parser_name: str,
    parser_status: str,
    parser_status_reason: str | None,
    chunks: list[ParsedDocumentChunk]
) -> ParsedDocument:
    document_version = str(metadata.get('document_version') or 'v1')
    revision_id = str(metadata.get('revision_id') or '')
    source_id = str(metadata.get('source_id') or f'{parser_name}:unknown')
    source_url = str(metadata.get('source_url') or source_id)
    content_signature = str(metadata.get('content_signature') or f'{source_id}:{document_version}')
    # Auto-generate source_snippet from the first chunk if not provided
    source_snippet = str(
        metadata.get('source_snippet')
        or (chunks[0].source_snippet if chunks else '')
    )

    return ParsedDocument(
        source_id=source_id,
        source_url=source_url,
        source_snippet=source_snippet,
        permission_level=str(metadata.get('permission_level') or 'internal'),
        mime_type=str(metadata.get('mime_type') or ''),
        document_version=document_version,
        revision_id=revision_id,
        content_signature=content_signature,
        parser_run=ParserRun(
            parser_name=parser_name,
            parser_status=parser_status,
            parser_status_reason=parser_status_reason,
        ),
        chunks=chunks,
    )

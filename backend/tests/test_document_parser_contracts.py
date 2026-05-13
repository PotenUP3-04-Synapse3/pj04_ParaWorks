import pytest

from backend.app.documents.parsers import (
    DocumentParserError,
    ParsedDocument,
    ParsedDocumentChunk,
    ParserRun,
    parser_adapter_decision_for_mime_type,
)


def test_parsed_document_requires_source_evidence_for_chunks() -> None:
    with pytest.raises(DocumentParserError, match='source evidence'):
        ParsedDocument(
            source_id='',
            source_url='https://drive.google.com/file/d/file-1/view',
            source_snippet='휴가 신청은 HR 시스템에서 진행합니다.',
            permission_level='restricted',
            mime_type='application/vnd.google-apps.document',
            document_version='42',
            revision_id='rev-42',
            content_signature='drive:file-1:42:rev-42',
            parser_run=ParserRun(
                parser_name='google_drive_text_export',
                parser_status='parsed',
                parser_status_reason=None,
            ),
            chunks=[
                ParsedDocumentChunk(
                    chunk_index=0,
                    text='휴가 신청은 HR 시스템에서 진행합니다.',
                    source_snippet='휴가 신청은 HR 시스템에서 진행합니다.',
                    section_path='휴가 정책',
                )
            ],
        )


def test_parsed_document_enriches_chunks_with_parser_metadata() -> None:
    parsed = ParsedDocument(
        source_id='drive:file-1',
        source_url='https://drive.google.com/file/d/file-1/view',
        source_snippet='휴가 신청은 HR 시스템에서 진행합니다.',
        permission_level='restricted',
        mime_type='application/vnd.google-apps.document',
        document_version='42',
        revision_id='rev-42',
        content_signature='drive:file-1:42:rev-42',
        parser_run=ParserRun(
            parser_name='google_drive_text_export',
            parser_status='parsed',
            parser_status_reason=None,
        ),
        chunks=[
            ParsedDocumentChunk(
                chunk_index=0,
                text='휴가 신청은 HR 시스템에서 진행합니다.',
                source_snippet='휴가 신청은 HR 시스템에서 진행합니다.',
                section_path='휴가 정책',
            )
        ],
    )

    chunk = parsed.chunks[0]
    assert chunk.permission_level == 'restricted'
    assert chunk.metadata['source_id'] == 'drive:file-1'
    assert chunk.metadata['source_url'] == 'https://drive.google.com/file/d/file-1/view'
    assert chunk.metadata['parser_name'] == 'google_drive_text_export'
    assert chunk.metadata['parser_status'] == 'parsed'
    assert chunk.metadata['document_version'] == '42'
    assert chunk.metadata['revision_id'] == 'rev-42'
    assert chunk.metadata['content_signature'] == 'drive:file-1:42:rev-42'
    assert len(chunk.content_hash) == 64


def test_parser_adapter_decision_documents_pdf_and_docx_candidates() -> None:
    pdf = parser_adapter_decision_for_mime_type('application/pdf')
    docx = parser_adapter_decision_for_mime_type(
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

    assert pdf.parser_status == 'parsed'
    assert pdf.parser_status_reason == ''
    assert pdf.candidate_package == 'pypdf'
    assert pdf.live_enabled is True
    assert docx.parser_status == 'parsed'
    assert docx.parser_status_reason == ''
    assert docx.candidate_package == 'python-docx'
    assert docx.live_enabled is True


def test_parser_adapter_decision_keeps_hwp_and_hwpx_unsupported() -> None:
    for mime_type in [
        'application/x-hwp',
        'application/haansofthwp',
        'application/vnd.hancom.hwpx',
    ]:
        decision = parser_adapter_decision_for_mime_type(mime_type)
        assert decision.parser_status == 'unsupported'
        assert decision.parser_status_reason == 'hwp_parser_not_decided'
        assert decision.candidate_package is None
        assert decision.live_enabled is False


def test_pdf_document_parser_extracts_text_by_page() -> None:
    import io
    from reportlab.pdfgen import canvas as rl_canvas  # type: ignore[import]
    from backend.app.documents.adapters import PdfDocumentParser

    buf = io.BytesIO()
    c = rl_canvas.Canvas(buf)
    c.drawString(100, 750, 'Hello PDF page one')
    c.showPage()
    c.drawString(100, 750, 'Hello PDF page two')
    c.save()
    buf.seek(0)
    payload = buf.read()

    metadata = {
        'source_id': 'drive:pdf-1',
        'source_url': 'https://drive.google.com/file/d/pdf-1/view',
        'permission_level': 'internal',
        'mime_type': 'application/pdf',
        'document_version': 'v1',
    }
    parsed = PdfDocumentParser().parse(payload, metadata=metadata)

    assert parsed.parser_run.parser_name == 'pypdf'
    assert parsed.parser_run.parser_status == 'parsed'
    assert len(parsed.chunks) == 2
    assert 'page one' in parsed.chunks[0].text.lower() or 'Hello PDF' in parsed.chunks[0].text
    assert parsed.chunks[0].page_number == 1
    assert parsed.chunks[1].page_number == 2


def test_docx_document_parser_extracts_paragraphs() -> None:
    import io
    from docx import Document as DocxDoc
    from backend.app.documents.adapters import DocxDocumentParser

    buf = io.BytesIO()
    doc = DocxDoc()
    doc.add_paragraph('첫 번째 단락입니다.')
    doc.add_paragraph('두 번째 단락입니다.')
    doc.save(buf)
    buf.seek(0)
    payload = buf.read()

    metadata = {
        'source_id': 'drive:docx-1',
        'source_url': 'https://drive.google.com/file/d/docx-1/view',
        'permission_level': 'internal',
        'mime_type': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        'document_version': 'v1',
    }
    parsed = DocxDocumentParser().parse(payload, metadata=metadata)

    assert parsed.parser_run.parser_name == 'python-docx'
    assert parsed.parser_run.parser_status == 'parsed'
    assert len(parsed.chunks) >= 1
    combined = '\n'.join(c.text for c in parsed.chunks)
    assert '첫 번째 단락입니다.' in combined
    assert '두 번째 단락입니다.' in combined


def test_pdf_document_parser_returns_error_status_on_invalid_payload() -> None:
    from backend.app.documents.adapters import PdfDocumentParser

    parsed = PdfDocumentParser().parse(b'not a pdf', metadata={})

    assert parsed.parser_run.parser_status == 'error'
    assert parsed.parser_run.parser_status_reason is not None
    assert 'Failed to parse PDF' in parsed.parser_run.parser_status_reason
    assert parsed.chunks == []


def test_text_document_parser_extracts_paragraphs() -> None:
    from backend.app.documents.adapters import TextDocumentParser

    payload = "First paragraph.\n\nSecond paragraph.".encode('utf-8')
    metadata = {
        'source_id': 'drive:text-1',
        'source_url': 'https://drive.google.com/file/d/text-1/view',
        'permission_level': 'internal',
        'mime_type': 'text/plain',
        'document_version': 'v1',
    }
    parsed = TextDocumentParser().parse(payload, metadata=metadata)

    assert parsed.parser_run.parser_name == 'built-in-text'
    assert parsed.parser_run.parser_status == 'parsed'
    # Both paragraphs are short enough to fit in one chunk
    assert len(parsed.chunks) == 1
    assert 'First paragraph.' in parsed.chunks[0].text
    assert 'Second paragraph.' in parsed.chunks[0].text


def test_markdown_document_parser_extracts_paragraphs() -> None:
    from backend.app.documents.adapters import TextDocumentParser

    payload = "# Title\n\nContent body.".encode('utf-8')
    metadata = {
        'source_id': 'drive:md-1',
        'source_url': 'https://drive.google.com/file/d/md-1/view',
        'permission_level': 'internal',
        'mime_type': 'text/markdown',
        'document_version': 'v1',
    }
    parsed = TextDocumentParser().parse(payload, metadata=metadata)

    assert parsed.parser_run.parser_name == 'built-in-text'
    assert parsed.parser_run.parser_status == 'parsed'
    assert len(parsed.chunks) == 1
    assert '# Title' in parsed.chunks[0].text
    assert 'Content body.' in parsed.chunks[0].text

